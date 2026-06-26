"""
Markdown → Word (.docx) 转换工具 —— 中文专业报告模板版

适用于:
    - 技术方案 / 项目交付文档
    - 运维手册 / 网络安全报告
    - 政企客户交付件 / 等保测评报告
    - DDoS分析报告 / SSO对接说明等

用法:
    python md2docx.py input.md                        # 输出为 input.docx
    python md2docx.py input.md -o output.docx         # 指定输出文件名
    python md2docx.py input.md --cover --header --footer --numbering --pdf

支持:
    - 标题 h1-h6（黑体，纯黑）
    - 粗体 / 斜体 / 粗斜体 / 删除线 / 行内代码（Bootstrap 红）
    - 链接（无下划线，蓝色）/ 图片（居中，自动编号）
    - 有序 / 无序 / 多级嵌套列表
    - 表格（智能列宽，单元格间距）
    - 代码块（单格表格包裹，灰底 + 边框）
    - 引用块（楷体 + 左侧色条 + 浅灰底）
    - 水平分隔线 / 脚注
    - 自动封面（--cover）+ 自动目录（TOC）
    - 页眉（--header）/ 页码（--footer）
    - 标题自动编号（--numbering，检测到手写编号时自动跳过）
    - PDF 输出（--pdf，需 WPS Office）
"""

import argparse
import os
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ═══════════════════════════════════════════════════════════
#  样式常量 —— 中文专业报告模板
# ═══════════════════════════════════════════════════════════

# 字体族
FONT_NAME_BODY_CN = "宋体"
FONT_NAME_TITLE_CN = "黑体"
FONT_NAME_SUBTITLE_CN = "黑体"
FONT_NAME_QUOTE_CN = "楷体"
FONT_NAME_TABLE_CN = "微软雅黑"
FONT_NAME_EN = "Times New Roman"
FONT_CODE = "Consolas"

# 字号
FONT_SIZE_BODY = Pt(12)       # 小四
FONT_SIZE_TABLE = Pt(10.5)    # 五号
FONT_SIZE_CODE = Pt(10.5)
FONT_SIZE_FOOTNOTE = Pt(9)

# 正文颜色
COLOR_BODY = RGBColor(0, 0, 0)                 # 纯黑
COLOR_HEADING_1 = "000000"
COLOR_HEADING_2 = "000000"
COLOR_HEADING_3 = "000000"
COLOR_HEADING_4 = "000000"
COLOR_LINK = "0563C1"
COLOR_CODE_LANG = "888888"

# 标题配置
HEADING_CONFIG = {
    1: {
        "size": Pt(22),
        "color": COLOR_HEADING_1,
        "space_before": Pt(30),
        "space_after": Pt(14),
    },
    2: {
        "size": Pt(18),
        "color": COLOR_HEADING_2,
        "space_before": Pt(22),
        "space_after": Pt(10),
    },
    3: {
        "size": Pt(15),
        "color": COLOR_HEADING_3,
        "space_before": Pt(14),
        "space_after": Pt(8),
    },
    4: {
        "size": Pt(13),
        "color": COLOR_HEADING_4,
        "space_before": Pt(10),
        "space_after": Pt(6),
    },
    5: {
        "size": Pt(12),
        "color": "808080",
        "space_before": Pt(6),
        "space_after": Pt(4),
    },
    6: {
        "size": Pt(10.5),
        "color": "808080",
        "space_before": Pt(6),
        "space_after": Pt(4),
    },
}

# 表格配色
TABLE_HEADER_BG = "1F4E79"
TABLE_ROW_ODD = "FFFFFF"
TABLE_ROW_EVEN = "F7FBFF"

# 引用块配色
BLOCKQUOTE_BG = "F0F4F8"
BLOCKQUOTE_BORDER = "4472C4"

# 代码块配色
CODE_BLOCK_BG = "F7F8FA"
CODE_BLOCK_BORDER = "D9D9D9"

# 图片最大宽度
MAX_IMAGE_WIDTH = Inches(6.2)

# ── 封面常量 ──
COVER_BG_COLOR = "1F4E79"
COVER_SUBTITLE = "技术方案说明书"

# 图片自动编号
_image_counter = [0]
# 表格自动编号
_table_counter = [0]

# 正文行距（国标固定值 28磅）
BODY_LINE_SPACING_FIXED = Pt(28)


# ═══════════════════════════════════════════════════════════
#  XML / OPC 辅助
# ═══════════════════════════════════════════════════════════


def set_cell_shading(cell, color_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_paragraph_shading(paragraph, color_hex):
    p_pr = paragraph._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def set_paragraph_left_border(paragraph, color_hex, width_pt=3):
    """给段落加左边框（用于 blockquote）"""
    p_pr = paragraph._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt * 8))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    borders.append(left)
    p_pr.append(borders)


def add_horizontal_line(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    p_pr.append(borders)


# ═══════════════════════════════════════════════════════════
#  页脚 — 页码
# ═══════════════════════════════════════════════════════════


def _add_page_number(section):
    """给 section 的页脚添加居中页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run("第 ")
    _set_run_font_simple(run, size=Pt(9), color="666666", font_name=FONT_NAME_BODY_CN)

    # PAGE 域
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run1 = p.add_run()
    run1._element.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2 = p.add_run()
    run2._element.append(instr)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run3 = p.add_run()
    run3._element.append(fld_char_sep)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run4 = p.add_run()
    run4._element.append(fld_char_end)

    run = p.add_run(" 页")
    _set_run_font_simple(run, size=Pt(9), color="666666", font_name=FONT_NAME_BODY_CN)


def _add_header(section, title="技术方案说明书"):
    """给 section 添加页眉（封面页不显示）"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    _set_run_font_simple(run, size=Pt(9), color="999999", font_name=FONT_NAME_BODY_CN)


# ═══════════════════════════════════════════════════════════
#  字体 / Run 辅助（细粒度函数族）
# ═══════════════════════════════════════════════════════════


def _set_run_font_simple(run, size=None, color=None, font_name=None):
    """基础的字体设置，不触及 eastAsia"""
    if font_name:
        run.font.name = font_name
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color


def _set_run_font_cn(run, cn_font, en_font=None, bold=False, italic=False,
                     strike=False, size=None, color=None):
    """同时设置中西文字体"""
    en_font = en_font or FONT_NAME_EN
    run.font.name = en_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    if size:
        run.font.size = size
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if strike:
        run.font.strike = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color


def set_body_font(run, bold=False, italic=False, size=None, **kwargs):
    """正文：宋体"""
    _set_run_font_cn(run, FONT_NAME_BODY_CN, bold=bold, italic=italic,
                     size=size or FONT_SIZE_BODY, color=COLOR_BODY)


def set_heading_font(run, level, size=None, **kwargs):
    """标题：黑体 + 对应颜色"""
    cfg = HEADING_CONFIG.get(level, HEADING_CONFIG[4])
    _set_run_font_cn(run, FONT_NAME_TITLE_CN, bold=True,
                     size=size or cfg["size"], color=cfg["color"])


def set_table_header_font(run, size=None, **kwargs):
    """表头：微软雅黑 + 粗体"""
    _set_run_font_cn(run, FONT_NAME_TABLE_CN, bold=True,
                     size=size or FONT_SIZE_TABLE, color="000000")


def set_table_cell_font(run, size=None, **kwargs):
    """表体：宋体"""
    _set_run_font_cn(run, FONT_NAME_BODY_CN, size=size or FONT_SIZE_TABLE,
                     color=COLOR_BODY)


def set_quote_font(run, size=None, **kwargs):
    """引用：楷体"""
    _set_run_font_cn(run, FONT_NAME_QUOTE_CN, italic=False,
                     size=size or FONT_SIZE_BODY, color=COLOR_BODY)


def set_code_font(run, size=None, **kwargs):
    """代码：Consolas"""
    _set_run_font_simple(run, font_name=FONT_CODE,
                         size=size or FONT_SIZE_CODE)


def set_code_inline_font(run, size=None):
    """行内代码：Consolas + Bootstrap 风格"""
    _set_run_font_simple(run, font_name=FONT_CODE,
                         size=size or FONT_SIZE_BODY, color="C7254E")


def set_link_font(run, size=None):
    """链接：无下划线，蓝色 0563C1"""
    _set_run_font_cn(run, FONT_NAME_BODY_CN, size=size or FONT_SIZE_BODY,
                     color=COLOR_LINK)
    run.underline = False


# ═══════════════════════════════════════════════════════════
#  内联标记解析
# ═══════════════════════════════════════════════════════════


def add_inline_runs(paragraph, text, base_size=None, base_bold=False,
                    base_italic=False, footnotes=None, context="body"):
    """
    将 Markdown 内联标记解析为 Word Run 对象。
    context 决定默认字体风格: body | table | quote | heading
    """
    if base_size is None:
        base_size = FONT_SIZE_BODY
    if footnotes is None:
        footnotes = {}

    # 根据 context 选择默认字体设置函数
    _default_font = {
        "body": lambda r, **kw: set_body_font(r, **kw),
        "table": lambda r, **kw: set_table_cell_font(r, **kw),
        "quote": lambda r, **kw: set_quote_font(r, **kw),
        "heading": lambda r, **kw: set_heading_font(r, kw.get("level", 4), size=kw.get("size")),
    }.get(context, set_body_font)

    combined = re.compile(
        r"(!\[[^\]]*\]\([^)]+\))"
        r"|(\[[^\]]+\]\([^)]+\))"
        r"|(\*\*\*.+?\*\*\*)"
        r"|(\*\*.+?\*\*)"
        r"|(\*(?!\*).+?(?<!\*)\*)"
        r"|(~~.+?~~)"
        r"|(`[^`]+`)"
        r"|(\[\^\w+\])"
    )

    pos = 0
    for m in combined.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos: m.start()])
            _default_font(run, bold=base_bold, italic=base_italic, size=base_size)

        matched = m.group(0)

        if matched.startswith("!["):
            img_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", matched)
            alt, src = img_m.group(1), img_m.group(2)  # type: ignore[union-attr]
            _try_insert_image(paragraph, src, alt)

        elif matched.startswith("[^"):
            fn_id = matched[2:-1]
            run = paragraph.add_run(f"[{fn_id}]")
            _set_run_font_simple(run, size=FONT_SIZE_FOOTNOTE, color="666666")

        elif matched.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", matched)
            link_text = link_m.group(1)  # type: ignore[union-attr]
            run = paragraph.add_run(link_text)
            set_link_font(run, size=base_size)

        elif matched.startswith("***"):
            inner = matched[3:-3]
            run = paragraph.add_run(inner)
            _default_font(run, bold=True, italic=True, size=base_size)

        elif matched.startswith("**"):
            inner = matched[2:-2]
            run = paragraph.add_run(inner)
            _default_font(run, bold=True, italic=base_italic, size=base_size)

        elif matched.startswith("~~"):
            inner = matched[2:-2]
            run = paragraph.add_run(inner)
            _default_font(run, strike=True, bold=base_bold, size=base_size)

        elif matched.startswith("`"):
            inner = matched[1:-1]
            run = paragraph.add_run(inner)
            set_code_inline_font(run, size=base_size)

        elif matched.startswith("*"):
            inner = matched[1:-1]
            run = paragraph.add_run(inner)
            _default_font(run, italic=True, bold=base_bold, size=base_size)

        pos = m.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _default_font(run, bold=base_bold, italic=base_italic, size=base_size)


def _try_insert_image(paragraph, src, alt):
    """尝试插入本地图片（居中 + 自动编号），失败则显示占位文本"""
    src_path = Path(src)
    if src_path.is_file():
        try:
            _image_counter[0] += 1
            # 图片居中
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(src_path), width=MAX_IMAGE_WIDTH)
            # 图注
            caption_p = paragraph.insert_paragraph_after()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = caption_p.add_run(f"图{_image_counter[0]} {alt}")
            _set_run_font_simple(cap_run, size=Pt(9), color="666666",
                                 font_name=FONT_NAME_BODY_CN)
            return
        except Exception:
            pass
    run = paragraph.add_run(f"[图片: {alt or src}]")
    set_body_font(run, italic=True, size=FONT_SIZE_BODY)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ═══════════════════════════════════════════════════════════
#  自动封面
# ═══════════════════════════════════════════════════════════


def add_cover_page(doc, title):
    """生成封面页"""
    # 用空段落占位，把标题推到页面中部
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        _set_run_font_simple(run, size=Pt(12))

    # 蓝色装饰条
    p_bar = doc.add_paragraph()
    p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bar.paragraph_format.space_before = Pt(24)
    p_bar.paragraph_format.space_after = Pt(12)
    run_bar = p_bar.add_run("━" * 30)
    _set_run_font_simple(run_bar, size=Pt(14), color=COVER_BG_COLOR)

    # 主标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    run_title = p_title.add_run(title)
    _set_run_font_cn(run_title, FONT_NAME_TITLE_CN, bold=True,
                     size=Pt(28), color=COVER_BG_COLOR)

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run(COVER_SUBTITLE)
    _set_run_font_cn(run_sub, FONT_NAME_SUBTITLE_CN, bold=False,
                     size=Pt(16), color="808080")

    # 蓝色装饰条
    p_bar2 = doc.add_paragraph()
    p_bar2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bar2.paragraph_format.space_before = Pt(12)
    p_bar2.paragraph_format.space_after = Pt(48)
    run_bar2 = p_bar2.add_run("━" * 30)
    _set_run_font_simple(run_bar2, size=Pt(14), color=COVER_BG_COLOR)

    # 日期
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today()
    date_str = f"{today.year}年{today.month}月{today.day}日"
    run_date = p_date.add_run(date_str)
    _set_run_font_simple(run_date, size=Pt(12), color="999999",
                         font_name=FONT_NAME_BODY_CN)

    _add_cover_footer(doc)

    # 封面后的分页
    doc.add_page_break()


def _add_cover_footer(doc):
    """封面底部免责/公司信息"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("— 本文档仅供内部使用 —")
    _set_run_font_simple(run, size=Pt(9), color="BBBBBB", font_name=FONT_NAME_BODY_CN)


# ═══════════════════════════════════════════════════════════
#  自动目录
# ═══════════════════════════════════════════════════════════


def add_toc(doc):
    """插入 Word TOC 域 (目录)"""
    p_heading = doc.add_paragraph()
    run = p_heading.add_run("目  录")
    _set_run_font_cn(run, FONT_NAME_TITLE_CN, bold=True,
                     size=Pt(16), color=COLOR_HEADING_1)
    p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_heading.paragraph_format.space_after = Pt(12)

    # TOC 域
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 使用 XML 插入 TOC 域
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run0 = p.add_run()
    run0._element.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run1 = p.add_run()
    run1._element.append(instr)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run2 = p.add_run()
    run2._element.append(fld_char_sep)

    # 占位文本
    run3 = p.add_run("（请在 Word 中右键点击此处 → 更新域 生成目录）")
    _set_run_font_simple(run3, size=Pt(10), color="999999",
                         font_name=FONT_NAME_BODY_CN)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run4 = p.add_run()
    run4._element.append(fld_char_end)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
#  元素构建
# ═══════════════════════════════════════════════════════════

_heading_numbering_setup_done = False
_heading_numbering_enabled = False


def _detect_existing_numbering(blocks):
    """扫描所有标题，检测是否已有手动编号"""
    for block in blocks:
        if block[0] == "heading":
            text = block[2]
            # 匹配常见的标题编号模式
            if re.match(
                r'^('
                r'\d+\b|'                      # 1 2 3
                r'\d+[\.、．)]\s*|'             # 1. 1、1) 1.
                r'\d+\.\d+\b|'                 # 1.1 2.3
                r'\d+\.\d+\.\d+\b|'            # 1.1.1
                r'[一二三四五六七八九十百千]+[、．]|'  # 一、 二、
                r'（[一二三四五六七八九十百千]+）|'   # （一）（二）
                r'第[一二三四五六七八九十百千]+[章节篇]|'  # 第一章
                r'[A-E]\)|'                    # A) B)
                r'Part\s+\d+|'                 # Part 1
                r'Chapter\s+\d+'               # Chapter 1
                r')\s',
                text
            ):
                return True
    return False


def _setup_heading_numbering(doc):
    """创建标题多级编号定义（1 / 1.1 / 1.1.1 ...）"""
    global _heading_numbering_setup_done
    if _heading_numbering_setup_done or not _heading_numbering_enabled:
        return
    _heading_numbering_setup_done = True

    numbering_part = doc.part.numbering_part
    numbering = numbering_part._element

    # 检查是否已存在
    for existing in numbering.findall(qn('w:abstractNum')):
        if existing.get(qn('w:abstractNumId')) == '100':
            return

    # Abstract numbering
    abs_num = OxmlElement('w:abstractNum')
    abs_num.set(qn('w:abstractNumId'), '100')

    # 四级标题编号格式
    formats = [
        {'fmt': 'decimal', 'text': '%1、'},
        {'fmt': 'decimal', 'text': '%1.%2'},
        {'fmt': 'decimal', 'text': '%1.%2.%3'},
        {'fmt': 'decimal', 'text': '%1.%2.%3.%4'},
    ]

    for level in range(4):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), formats[level]['fmt'])
        lvl.append(numFmt)

        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), formats[level]['text'])
        lvl.append(lvlText)

        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), f'heading{level + 1}')
        lvl.append(pStyle)

        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # 编号字体 + 颜色（与对应标题一致）
        rPr = OxmlElement('w:rPr')
        cfg = HEADING_CONFIG.get(level + 1, HEADING_CONFIG[4])
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE_CN)
        rFonts.set(qn('w:ascii'), FONT_NAME_EN)
        rFonts.set(qn('w:hAnsi'), FONT_NAME_EN)
        rPr.append(rFonts)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), cfg['color'])
        rPr.append(color)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(cfg['size'].pt * 2)))
        rPr.append(sz)

        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(cfg['size'].pt * 2)))
        rPr.append(szCs)

        lvl.append(rPr)
        abs_num.append(lvl)

    numbering.append(abs_num)

    # Numbering instance
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '100')
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), '100')
    num.append(abstractNumId)
    numbering.append(num)


def _add_numpr_to_heading(paragraph, level):
    """给标题段落添加编号引用"""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '100')
    numPr.append(numId)

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level - 1))
    numPr.append(ilvl)

    pPr.append(numPr)


def add_heading(doc, text, level, footnotes=None):
    cfg = HEADING_CONFIG.get(level, HEADING_CONFIG[4])
    h = doc.add_heading(level=min(level, 4))
    # 清除默认 run
    for r in h.runs:
        r.text = ""
    h.paragraph_format.space_before = cfg["space_before"]
    h.paragraph_format.space_after = cfg["space_after"]
    # 一级标题居中
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 仅当未检测到手动编号时才添加自动编号
    if _heading_numbering_enabled:
        _add_numpr_to_heading(h, level)
    add_inline_runs(h, text, base_size=cfg["size"], base_bold=True,
                    footnotes=footnotes, context="heading")


def add_body_paragraph(doc, text, footnotes=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = BODY_LINE_SPACING_FIXED
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
    add_inline_runs(p, text, footnotes=footnotes, context="body")
    return p


def _content_width(text):
    """估算文本宽度：中文算2个单位，英文/数字算1个"""
    w = 0
    for ch in text:
        w += 2 if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' else 1
    return w


def _set_column_widths(table, headers, rows):
    """按内容比例分配列宽"""
    col_count = len(headers)
    # 计算每列最大宽度
    col_widths = [_content_width(h) for h in headers]
    for row in rows:
        for i in range(min(len(row), col_count)):
            col_widths[i] = max(col_widths[i], _content_width(row[i]))
    # 加一点余量
    col_widths = [max(w + 2, 6) for w in col_widths]
    total = sum(col_widths)
    # A4 正文可用宽度 ≈ 14.66cm (21 - 3.17*2)
    available = Cm(14.66)
    for i in range(col_count):
        ratio = col_widths[i] / total
        for row in table.rows:
            row.cells[i].width = int(available * ratio)


def add_table(doc, headers, rows):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    # 智能列宽：按内容长度分配
    _set_column_widths(table, headers, rows)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_table_header_font(run)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx + 1].cells[c_idx]
            val = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            add_inline_runs(cell.paragraphs[0], val, base_size=FONT_SIZE_TABLE,
                            context="table")
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # 表号 + 表注（取第一列标题作为表名）
    _table_counter[0] += 1
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(8)
    caption_text = f"表{_table_counter[0]} {headers[0] if headers else ''}"
    run = cap_p.add_run(caption_text)
    _set_run_font_simple(run, size=Pt(9), color="666666", font_name=FONT_NAME_BODY_CN)


def add_code_block(doc, code_text, language=""):
    """添加代码块：单格表格包裹，整体灰底 + 圆角感边框"""
    # 一行搞定：用单格表格把所有代码包在一起
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 移除表格默认样式
    table.style = "Table Grid"
    # 移除所有内边距边距
    cell = table.rows[0].cells[0]
    # 清除默认段落
    cell.paragraphs[0].clear()

    # 语言标签（如果代码块有标注语言）
    if language:
        p_lang = cell.paragraphs[0]
        run = p_lang.add_run(f"  {language}")
        _set_run_font_simple(run, size=Pt(8), color=COLOR_CODE_LANG,
                             font_name=FONT_CODE)
        p_lang.paragraph_format.space_after = Pt(2)
        p_lang.paragraph_format.space_before = Pt(2)

    # 代码内容（单一段落，用换行符分隔）
    # 将代码塞入一个段落，保持行结构
    p_code = cell.add_paragraph() if language else cell.paragraphs[0]
    p_code.paragraph_format.line_spacing = 1.2
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(4)

    lines = code_text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            p_code.add_run("\n")
        if line:
            run = p_code.add_run(line)
            set_code_font(run)
        else:
            # 空行：插入零宽度的换行占位
            run = p_code.add_run("\u200B")
            set_code_font(run)

    # 设置单元格底色
    set_cell_shading(cell, CODE_BLOCK_BG)
    # 设置单元格四周边框
    _set_code_cell_border(cell)

    # 缩小单元格边距
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", "40"), ("left", "80"), ("bottom", "40"), ("right", "80")]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    doc.add_paragraph()  # 代码块后空行


def _set_code_cell_border(cell):
    """给代码块单元格设置四周边框"""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), CODE_BLOCK_BORDER)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def add_blockquote(doc, lines, footnotes=None):
    """添加引用块：楷体 + 左侧色条 + 浅灰底"""
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = BODY_LINE_SPACING_FIXED
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(p, BLOCKQUOTE_BG)
        set_paragraph_left_border(p, BLOCKQUOTE_BORDER, width_pt=2)
        text = line.lstrip("> ").strip()
        add_inline_runs(p, text, base_size=FONT_SIZE_BODY, base_italic=False,
                        footnotes=footnotes, context="quote")


def add_list_item(doc, text, level=0, ordered=False, number=1, footnotes=None):
    """添加列表项（支持嵌套）"""
    if ordered:
        p = doc.add_paragraph()
        indent = Cm(1.27 * (level + 1))
        p.paragraph_format.left_indent = indent
        p.paragraph_format.first_line_indent = Cm(-0.63)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{number}. ")
        set_body_font(run, size=FONT_SIZE_BODY)
        add_inline_runs(p, text, footnotes=footnotes, context="body")
    else:
        style_name = "List Bullet"
        if level == 1:
            style_name = "List Bullet 2"
        elif level >= 2:
            style_name = "List Bullet 3"
        try:
            p = doc.add_paragraph(style=style_name)
        except KeyError:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline_runs(p, text, footnotes=footnotes, context="body")


# ═══════════════════════════════════════════════════════════
#  Markdown 解析器
# ═══════════════════════════════════════════════════════════


def _detect_list_indent(line):
    stripped = line.rstrip()
    spaces = len(stripped) - len(stripped.lstrip())
    level = spaces // 2
    return level


def parse_md(md_text):
    blocks = []
    footnotes = {}
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # HTML 注释
        if stripped.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # 脚注定义
        fn_m = re.match(r"^\[\^(\w+)\]:\s*(.*)", stripped)
        if fn_m:
            fn_id, fn_text = fn_m.group(1), fn_m.group(2)
            i += 1
            while i < len(lines) and lines[i].startswith("  "):
                fn_text += " " + lines[i].strip()
                i += 1
            footnotes[fn_id] = fn_text
            continue

        # 代码块
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines), language))
            i += 1
            continue

        # 标题
        h_m = re.match(r"^(#{1,6})\s+(.*)", line)
        if h_m:
            level = len(h_m.group(1))
            blocks.append(("heading", level, h_m.group(2).strip()))
            i += 1
            continue

        # Setext 标题
        if i + 1 < len(lines) and stripped and re.match(r"^[=]{3,}$", lines[i + 1].strip()):
            blocks.append(("heading", 1, stripped))
            i += 2
            continue
        if i + 1 < len(lines) and stripped and re.match(r"^[-]{3,}$", lines[i + 1].strip()) and not re.match(r"^[-*]\s", stripped):
            blocks.append(("heading", 2, stripped))
            i += 2
            continue

        # 表格
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append(("table", headers, rows))
            continue

        # 水平线
        if re.match(r"^(\*{3,}|-{3,}|_{3,})$", stripped):
            blocks.append(("hr",))
            i += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip())
                i += 1
            blocks.append(("blockquote", quote_lines))
            continue

        # 有序列表
        ol_m = re.match(r"^(\s*)(\d+)[.)]\s+(.*)", line)
        if ol_m:
            list_items = []
            while i < len(lines):
                ol_line = re.match(r"^(\s*)(\d+)[.)]\s+(.*)", lines[i])
                if ol_line:
                    level = _detect_list_indent(lines[i])
                    num = int(ol_line.group(2))
                    list_items.append((level, num, ol_line.group(3).strip()))
                    i += 1
                elif i < len(lines) and lines[i].startswith("  ") and list_items:
                    prev = list_items[-1]
                    list_items[-1] = (prev[0], prev[1], prev[2] + " " + lines[i].strip())
                    i += 1
                else:
                    break
            blocks.append(("ordered_list", list_items))
            continue

        # 无序列表
        ul_m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if ul_m:
            list_items = []
            while i < len(lines):
                ul_line = re.match(r"^(\s*)[-*+]\s+(.*)", lines[i])
                if ul_line:
                    level = _detect_list_indent(lines[i])
                    list_items.append((level, ul_line.group(2).strip()))
                    i += 1
                elif i < len(lines) and lines[i].startswith("  ") and list_items:
                    prev = list_items[-1]
                    list_items[-1] = (prev[0], prev[1] + " " + lines[i].strip())
                    i += 1
                else:
                    break
            blocks.append(("unordered_list", list_items))
            continue

        # 普通文本
        if stripped:
            para_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip() \
                    and not lines[i].strip().startswith("#") \
                    and not lines[i].strip().startswith(">") \
                    and not lines[i].strip().startswith("```") \
                    and not lines[i].strip().startswith("|") \
                    and not re.match(r"^(\*{3,}|-{3,}|_{3,})$", lines[i].strip()) \
                    and not re.match(r"^[-*+]\s", lines[i].strip()) \
                    and not re.match(r"^\d+[.)]\s", lines[i].strip()) \
                    and not re.match(r"^\[\^", lines[i].strip()):
                para_lines.append(lines[i].strip())
                i += 1
            blocks.append(("text", " ".join(para_lines)))
            continue

        i += 1

    return blocks, footnotes


# ═══════════════════════════════════════════════════════════
#  文档构建
# ═══════════════════════════════════════════════════════════


def build_docx(blocks, footnotes, with_cover=False, with_numbering=False,
               with_header=False, with_footer=False):
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 获取文档标题（第一个 h1）
    doc_title = None
    for b in blocks:
        if b[0] == "heading" and b[1] == 1:
            doc_title = b[2]
            break

    # ── 自动封面（可选，--cover 参数开启）──
    if with_cover and doc_title:
        # 封面页使用独立的第一页页眉
        section.different_first_page_header_footer = True
        add_cover_page(doc, doc_title)
        add_toc(doc)

    # ── 页眉（可选，--header 参数开启）──
    if with_header:
        _add_header(section, title=doc_title or COVER_SUBTITLE)

    # ── 页脚页码（可选，--footer 参数开启）──
    if with_footer:
        _add_page_number(section)

    # ── 标题自动编号（仅当 --numbering 开启，且原文无手动编号时）──
    global _heading_numbering_enabled
    _heading_numbering_enabled = with_numbering and not _detect_existing_numbering(blocks)
    _setup_heading_numbering(doc)

    # ── 正文 ──
    for block in blocks:
        btype = block[0]

        if btype == "heading":
            _, level, text = block
            add_heading(doc, text, level, footnotes)

        elif btype == "text":
            add_body_paragraph(doc, block[1], footnotes)

        elif btype == "unordered_list":
            for level, text in block[1]:
                add_list_item(doc, text, level=level, ordered=False, footnotes=footnotes)

        elif btype == "ordered_list":
            for level, num, text in block[1]:
                add_list_item(doc, text, level=level, ordered=True, number=num, footnotes=footnotes)

        elif btype == "table":
            _, headers, rows = block
            add_table(doc, headers, rows)

        elif btype == "code":
            _, code_text, language = block
            add_code_block(doc, code_text, language)

        elif btype == "blockquote":
            add_blockquote(doc, block[1], footnotes)

        elif btype == "hr":
            add_horizontal_line(doc)

    # ── 脚注附录 ──
    if footnotes:
        doc.add_paragraph()
        add_horizontal_line(doc)
        h = doc.add_heading(level=3)
        p_pr = h._element.get_or_add_pPr()
        p_spacing = p_pr.find(qn("w:spacing"))
        if p_spacing is None:
            p_spacing = OxmlElement("w:spacing")
            p_pr.append(p_spacing)
        run = h.add_run("注释")
        _set_run_font_cn(run, FONT_NAME_TITLE_CN, bold=True,
                         size=Pt(14), color=COLOR_HEADING_3)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        for fn_id, fn_text in footnotes.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"[{fn_id}] ")
            _set_run_font_simple(run, bold=True, size=FONT_SIZE_FOOTNOTE,
                                 color="666666")
            add_inline_runs(p, fn_text, base_size=FONT_SIZE_FOOTNOTE, context="body")

    return doc


# ═══════════════════════════════════════════════════════════
#  CLI 入口
# ═══════════════════════════════════════════════════════════


def main():
    parser = argparse.ArgumentParser(
        description="Markdown → Word (.docx) 转换工具 —— 中文专业报告模板",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("input", help="输入 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径（默认与输入同名 .docx）")
    parser.add_argument(
        "--cover", action="store_true",
        help="生成自动封面 + 目录（默认不生成）"
    )
    parser.add_argument(
        "--numbering", action="store_true",
        help="标题自动编号（检测到原文已有编号时自动跳过）"
    )
    parser.add_argument(
        "--header", action="store_true",
        help="添加页眉（文档标题）"
    )
    parser.add_argument(
        "--footer", action="store_true",
        help="添加页脚（第 X 页）"
    )
    parser.add_argument(
        "--pdf", action="store_true",
        help="同时生成 PDF（需本机安装 WPS Office）"
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.is_file():
        print(f"错误: 找不到文件 {input_path}")
        return

    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    md_text = input_path.read_text(encoding="utf-8")
    blocks, footnotes = parse_md(md_text)
    doc = build_docx(blocks, footnotes, with_cover=args.cover,
                     with_numbering=args.numbering,
                     with_header=args.header, with_footer=args.footer)
    doc.save(str(output_path))
    print(f"✓ 已生成: {output_path}")

    # ── 可选：转 PDF（通过 WPS COM 接口）──
    if args.pdf:
        _convert_to_pdf(str(output_path))


def _convert_to_pdf(docx_path):
    """调用 WPS 将 docx 转为 pdf"""
    pdf_path = str(Path(docx_path).with_suffix(".pdf"))
    try:
        import win32com.client
        app = win32com.client.Dispatch("Kwps.Application")
        app.Visible = False
        doc = app.Documents.Open(os.path.abspath(docx_path))
        doc.ExportAsFixedFormat(os.path.abspath(pdf_path), 17)  # 17 = wdExportFormatPDF
        doc.Close()
        app.Quit()
        print(f"✓ 已生成: {pdf_path}")
    except Exception as e:
        print(f"⚠ PDF 转换失败: {e}")
        print("  请确认已安装 WPS Office 且版本支持 COM 自动化")


if __name__ == "__main__":
    main()
